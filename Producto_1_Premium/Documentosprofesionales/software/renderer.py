import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import List

try:
    # Importación relativa para ejecución como paquete
    from ..backend.schemas import FullDocumentSchema, ParagraphBlock
except ImportError:
    # Fallback para pruebas o ejecución directa
    import sys
    sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
    from backend.schemas import FullDocumentSchema, ParagraphBlock

class APADocBuilder:
    def __init__(self, output_filename="Paper_APA_Final.docx"):
        self.doc = Document()
        self.output_filename = output_filename
        self._configure_styles()

    def _configure_styles(self):
        """Configura la tipografía base (Times New Roman 12) y márgenes."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Configuración de márgenes (1 pulgada)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_page_number(self):
        """Hack avanzado para insertar números de página en el header (Campo XML 'PAGE')."""
        header = self.doc.sections[0].header
        # Aseguramos que haya un párrafo en el header
        if len(header.paragraphs) == 0:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
            
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        
        # Campo complejo XML para número de página
        fldChar1 = run._element.makeelement(qn('w:fldChar'))
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = run._element.makeelement(qn('w:instrText'))
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = run._element.makeelement(qn('w:fldChar'))
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

    def build_title_page(self, metadata):
        """Construye la portada oficial."""
        self._add_page_number()
        
        # Espaciado vertical inicial (aprox 3-4 líneas vacías)
        for _ in range(4): self.doc.add_paragraph()
        
        # Título en Negrita
        p_title = self.doc.add_paragraph(metadata.title)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.runs[0].bold = True
        
        # Espacio extra
        self.doc.add_paragraph()
        
        # Bloque de autoría
        details = [metadata.author, metadata.institution, "", f"Fecha: {metadata.date}"]
        for line in details:
            if line:
                p = self.doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        self.doc.add_page_break()

    def build_body(self, content_blocks: List[ParagraphBlock]):
        """Itera sobre los bloques validados e inyecta según su rol."""
        # Repetir título en negrita centrado al inicio del cuerpo (APA Std)
        # (Opcional según instrucciones, pero común en APA)
        
        for block in content_blocks:
            p = self.doc.add_paragraph()
            
            if block.role == "titulo1":
                p.text = block.content
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
            
            elif block.role == "titulo2":
                p = self.doc.add_paragraph(block.content)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.runs[0].bold = True
            
            elif block.role == "cuerpo":
                p = self.doc.add_paragraph(block.content)
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.line_spacing = 2.0
            
            elif block.role == "lista_item":
                # Lista con viñeta
                p = self.doc.add_paragraph(block.content, style='List Bullet')
                p.paragraph_format.line_spacing = 2.0
            
            elif block.role == "tabla":
                # Renderizado de Tabla
                table_data = block.content # Es un objeto TableBlock (o dict si viene de JSON crudo)
                
                # Normalización si viene como dict
                if isinstance(table_data, dict):
                    headers = table_data.get('headers', [])
                    rows = table_data.get('rows', [])
                else: 
                    headers = table_data.headers
                    rows = table_data.rows

                # Crear tabla
                cols = len(headers)
                table = self.doc.add_table(rows=1, cols=cols)
                table.style = 'Table Grid'
                
                # Header
                hdr_cells = table.rows[0].cells
                for i, h_text in enumerate(headers):
                    hdr_cells[i].text = str(h_text)
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                
                # Rows
                for row in rows:
                    if isinstance(row, dict): cells = row.get('cells', [])
                    else: cells = row.cells
                    
                    row_cells = table.add_row().cells
                    for i, cell_text in enumerate(cells):
                        if i < len(row_cells):
                            row_cells[i].text = str(cell_text)

            elif block.role == "cita_larga":
                p = self.doc.add_paragraph(block.content)
                p.paragraph_format.left_indent = Inches(0.5) # Sangría de bloque
                p.paragraph_format.line_spacing = 2.0
            
            # Fallback para roles no manejados o 'referencia' si viniera mezclado (no debería)

    def build_references(self, references_list: List[str]):
        """Crea la página de referencias con sangría francesa."""
        if not references_list: return
        
        self.doc.add_page_break()
        p_title = self.doc.add_paragraph("Referencias")
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.runs[0].bold = True
        
        for ref_text in references_list:
            p = self.doc.add_paragraph(ref_text)
            p.paragraph_format.line_spacing = 2.0
            # Sangría Francesa (Hanging Indent)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)

    def save(self):
        # Asegurar directorio de salida existe
        os.makedirs(os.path.dirname(self.output_filename), exist_ok=True)
        self.doc.save(self.output_filename)
        print(f"Documento generado exitosamente: {os.path.abspath(self.output_filename)}")

def run_pipeline(validated_data: FullDocumentSchema, output_path: str):
    """Interfaz de alto nivel para el pipeline."""
    builder = APADocBuilder(output_filename=output_path)
    builder.build_title_page(validated_data.metadata)
    builder.build_body(validated_data.content)
    builder.build_references(validated_data.references)
    builder.save()
