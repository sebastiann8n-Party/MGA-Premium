import os
from docx import Document

def analyze_docx(file_path):
    print(f"\n{'='*50}")
    print(f"ANALIZANDO: {os.path.basename(file_path)}")
    print(f"{'='*50}")
    
    if not os.path.exists(file_path):
        print(f"Archivo no encontrado: {file_path}")
        return

    try:
        doc = Document(file_path)
        
        # 1. Análisis de Secciones (Márgenes)
        print("\n--- [1. Configuración de Página] ---")
        for i, section in enumerate(doc.sections):
            top = section.top_margin.inches if section.top_margin else 0
            bottom = section.bottom_margin.inches if section.bottom_margin else 0
            left = section.left_margin.inches if section.left_margin else 0
            right = section.right_margin.inches if section.right_margin else 0
            print(f"Sección {i+1}: Márgenes (T:{top:.2f}, B:{bottom:.2f}, L:{left:.2f}, R:{right:.2f}) inches")
            
            # Header/Footer
            has_header = not section.header.is_linked_to_previous
            has_footer = not section.footer.is_linked_to_previous
            print(f"   Header propio: {has_header}, Footer propio: {has_footer}")

        # 2. Análisis de Estilos en Uso
        print("\n--- [2. Estilos Detectados en Párrafos] ---")
        styles_used = set()
        fonts_used = set()
        
        # Muestreo de los primeros 50 párrafos para no saturar
        for p in doc.paragraphs[:50]:
            if p.text.strip():
                styles_used.add(p.style.name)
                for run in p.runs:
                    if run.font.name:
                        fonts_used.add(run.font.name)
        
        print(f"Estilos únicos detectados: {list(styles_used)}")
        print(f"Fuentes explícitas detectadas: {list(fonts_used)}")

        # 3. Análisis de Tablas
        print(f"\n--- [3. Tablas] ---")
        print(f"Total de tablas: {len(doc.tables)}")
        if len(doc.tables) > 0:
            first_table = doc.tables[0]
            print(f"Estilo de primera tabla: '{first_table.style.name}'")
            print(f"Filas: {len(first_table.rows)}, Columnas: {len(first_table.columns)}")

        # 4. Texto de Muestra (Primeros 3 parrafos)
        print("\n--- [4. Muestra de Contenido] ---")
        for p in doc.paragraphs[:3]:
            if p.text.strip():
                print(f"[{p.style.name}] {p.text[:100]}...")

    except Exception as e:
        print(f"Error al leer DOCX: {e}")

if __name__ == "__main__":
    # Archivos de referencia
    ref_1 = r"x:\skills-analista\contexto\Producto_1_Premium\CITES_MGA_Proyecto.docx"
    ref_2 = r"x:\skills-analista\contexto\Producto_1_Premium\CITES_Apartado_Tecnico_MGA.docx"
    
    analyze_docx(ref_1)
    analyze_docx(ref_2)
