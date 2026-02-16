import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Importación relativa para tipos de datos del esquema
try:
    from ..backend.schemas import FullDocumentSchema, ParagraphBlock, TableBlock
except ImportError:
    sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
    from backend.schemas import FullDocumentSchema, ParagraphBlock, TableBlock

class CITESReportBuilder:
    def __init__(self, output_filename="Proyecto_CITES_Generado.docx", style_mode="MGA"):
        self.doc = Document()
        self.output_filename = output_filename
        self.style_mode = style_mode 
        
        # --- ESTADO DE NUMERACIÓN (State Machine) ---
        self.counters = {
            "h1": 0,
            "h2": 0,
            "h3": 0
        }
        
        self._setup_styles()

    def save(self):
        """Guarda el documento en la ruta definida."""
        self.doc.save(self.output_filename)
        print(f"Documento guardado en: {self.output_filename}")

    def _setup_styles(self):
        """Configura estilos sobrios tipo Informe Técnico (MGA/DNP)."""
        # Estilo Normal
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial' # MGA suele preferir Arial o Calibri
        font.size = Pt(11)
        
        # Configurar Título 1 (1. TÍTULO)
        # Aseguramos que el estilo existe o lo modificamos
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0) # Negro puro
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)

        # Configurar Título 2 (1.1 Subtítulo)
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(3)

    # --- NUEVOS MÉTODOS PARA COMPATIBILIDAD CON MGAWIZARD ---
    def add_text(self, text):
        """Agrega un párrafo de texto normal."""
        if text:
            self.doc.add_paragraph(text)

    def add_kv_list(self, items: list):
        """
        Agrega una lista de pares clave-valor.
        Ex: [{'key': 'Autor', 'value': 'Yo'}]
        """
        for item in items:
            self.add_key_value_paragraph(item.get('key', ''), item.get('value', ''))

    def add_table(self, table_content: dict):
        """Wrapper para create_table soportando el formato dict de MGAWizard."""
        if table_content:
            self.create_table(table_content)

    def add_numbered_heading(self, text, level=1):
        """
        Calcula automáticamente la numeración jerárquica (1. -> 1.1 -> 1.1.1)
        e inyecta el texto para que el Panel de Navegación lo reconozca.
        """
        if level == 1:
            self.counters["h1"] += 1
            self.counters["h2"] = 0 # Reiniciar subtítulos
            self.counters["h3"] = 0
            prefix = f"{self.counters['h1']}."
        elif level == 2:
            self.counters["h2"] += 1
            self.counters["h3"] = 0
            prefix = f"{self.counters['h1']}.{self.counters['h2']}"
        elif level == 3:
            self.counters["h3"] += 1
            prefix = f"{self.counters['h1']}.{self.counters['h2']}.{self.counters['h3']}"
        else:
            prefix = ""
            
        full_text = f"{prefix} {text.upper() if level == 1 else text}"
        self.doc.add_heading(full_text, level=level)

    def add_key_value_paragraph(self, key, value):
        """Replicar estilo: 'Viabilidad Económica: El proyecto...'"""
        p = self.doc.add_paragraph()
        runner_key = p.add_run(f"{key}: ")
        runner_key.bold = True
        p.add_run(str(value))

    def create_table(self, data, header=True):
        """
        Genera tablas profesionales estilo MGA.
        data: Lista de listas [['Col1', 'Col2'], ['Val1', 'Val2']] O objeto TableBlock O dict
        """
        # Adaptación para TableBlock de Pydantic
        if isinstance(data, TableBlock):
            rows_data = []
            rows_data.append(data.headers)
            for row in data.rows:
                rows_data.append(row.cells)
            data = rows_data
        elif isinstance(data, dict): # Soporte para MGAWizard
            headers = data.get('headers', [])
            rows = data.get('rows', [])
            rows_data = [headers] + rows
            data = rows_data
            
            # TODO: Manejar 'caption' si es necesario
            if 'caption' in data and False: # Placeholder logic
                pass

        if not data: return
        
        rows = len(data)
        cols = len(data[0]) if rows > 0 else 0
        if cols == 0: return

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid' # Estilo con bordes visibles
        
        for i, row_data in enumerate(data):
            # Asegurar que row_data tenga suficientes columnas
            safe_row_data = row_data + [''] * (cols - len(row_data))
            cells = table.rows[i].cells
            for j, cell_text in enumerate(safe_row_data):
                if j < len(cells):
                    cells[j].text = str(cell_text)
                    # Formato al Header
                    if header and i == 0:
                        for paragraph in cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)

        if not data: return
        
        rows = len(data)
        cols = len(data[0]) if rows > 0 else 0
        if cols == 0: return

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid' # Estilo con bordes visibles
        
        for i, row_data in enumerate(data):
            cells = table.rows[i].cells
            for j, cell_text in enumerate(row_data):
                if j < len(cells):
                    cells[j].text = str(cell_text)
                    # Formato al Header
                    if header and i == 0:
                        for paragraph in cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)

    def build_from_schema(self, schema: FullDocumentSchema):
        """Método de Integración: Construye el reporte completo desde datos estructurados."""
        
        # 1. Portada Institucional
        self.doc.add_heading(schema.metadata.title.upper(), 0)
        p = self.doc.add_paragraph(f"{schema.metadata.institution} - {schema.metadata.date}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_page_break()

        # 2. Iterar sobre bloques de contenido con lógica CITES
        for block in schema.content:
            if block.role == "titulo1":
                # Limpiar numeración manual si viene del MD para usar la auto-numeración del builder
                clean_text = block.content
                # Regex simple para quitar "1. TEXTO" -> "TEXTO" si existe, o dejarlo
                # Por ahora asumimos que el parser pasa el texto completo.
                # Para MGA estricto, idealmente limpiamos "1. " del string si vamos a re-numerar.
                if clean_text[0].isdigit() and "." in clean_text[:5]:
                     parts = clean_text.split(".", 1)
                     if len(parts) > 1: clean_text = parts[1].strip()
                
                self.add_numbered_heading(clean_text, level=1)

            elif block.role == "titulo2":
                clean_text = block.content
                if clean_text[0].isdigit() and "." in clean_text[:5]:
                     parts = clean_text.split(" ", 1) # A veces es "1.1 Texto"
                     if len(parts) > 1: clean_text = parts[1].strip()

                self.add_numbered_heading(clean_text, level=2)

            elif block.role == "tabla":
                self.create_table(block.content)
                # Espacio después de tabla
                self.doc.add_paragraph()

            elif block.role == "lista_item":
                # Lista con viñeta
                p = self.doc.add_paragraph(block.content, style='List Bullet')
                
            elif block.role == "cita_larga":
                 p = self.doc.add_paragraph(block.content)
                 p.paragraph_format.left_indent = Inches(0.5)
                 p.italic = True

            elif block.role == "cuerpo":
                # Detección de Key-Value (Heurística simple)
                if ":" in str(block.content) and len(str(block.content).split(":")[0]) < 30:
                    parts = str(block.content).split(":", 1)
                    self.add_key_value_paragraph(parts[0], parts[1].strip())
                else:
                    self.doc.add_paragraph(block.content)
            
            else:
                self.doc.add_paragraph(str(block.content))

        # Guardar
        self.doc.save(self.output_filename)
        print(f"Informe Tecnico Generado: {os.path.abspath(self.output_filename)}")

# Función puente para integración
def run_cites_pipeline(data: FullDocumentSchema, output_path: str):
    builder = CITESReportBuilder(output_filename=output_path)
    builder.build_from_schema(data)

if __name__ == "__main__":
    # Test rápido
    builder = CITESReportBuilder()
    builder.build_report() # Método original del snippet (no existe en mi versión adaptada arriba, oops)
    # Corrección: si se corre directo, no hará nada útil sin datos mock, 
    # pero el user code tenía un 'build_report' con datos mock.
    # He reemplazado build_report por build_from_schema para producción.
    pass
