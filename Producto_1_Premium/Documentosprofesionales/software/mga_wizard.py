import os
import sys

# Ajuste de path para importar módulos hermanos
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from software.cites_builder import CITESReportBuilder
from docx.enum.text import WD_ALIGN_PARAGRAPH

class MGAWizard:
    def __init__(self):
        self.data = {
            "project_title": "",
            "entidad": "",
            "sections": []
        }

    def ask(self, prompt):
        try:
            return input(f"{prompt}: ").strip()
        except EOFError:
            return ""

    def header(self, text):
        print(f"\n{'='*50}")
        print(f" {text}")
        print(f"{'='*50}")

    def run(self):
        self.header("GENERADOR DE PROYECTOS MGA (CLI MODE)")
        print("Este asistente le guiará para crear un documento técnico bajo estándares DNP.\n")

        # 1. Metadata
        print("--- PASO 1: INFORMACIÓN GENERAL ---")
        self.data["project_title"] = self.ask("Título del Proyecto")
        self.data["entidad"] = self.ask("Entidad Proponente")
        
        # 2. Identificación
        self.collect_identification()
        
        # 3. Técnica
        self.collect_technical()
        
        # 4. Presupuesto
        self.collect_financial()
        
        # 5. Riesgos
        self.collect_risks()
        
        # Generar
        self.generate()

    def collect_identification(self):
        self.header("PASO 2: IDENTIFICACIÓN DEL PROBLEMA")
        prob = self.ask("¿Cuál es el problema central a resolver?")
        pob = self.ask("Población Objetivo (ej. 500 Familias)")
        ubic = self.ask("Ubicación Principal (Municipio/Depto)")
        obj_gen = self.ask("Objetivo General del Proyecto")
        
        section = {
            "title": "Identificación y Objetivos",
            "level": 1,
            "blocks": [
                {"type": "texto", "text_content": f"Problema Central: {prob}"},
                {"type": "texto", "text_content": f"Objetivo General: {obj_gen}"},
                {"type": "kv_list", "kv_content": [
                    {"key": "Población Afectada", "value": pob},
                    {"key": "Ubicación Geográfica", "value": ubic}
                ]}
            ]
        }
        self.data["sections"].append(section)

    def collect_technical(self):
        self.header("PASO 3: ESTUDIO TÉCNICO")
        print("Describa la alternativa técnica seleccionada.")
        desc = self.ask("Descripción")
        
        items = []
        print("\n[Especificaciones Técnicas] (Presione Enter en 'Clave' para terminar)")
        while True:
            k = self.ask("  > Concepto/Clave (ej. Vida Útil)")
            if not k: break
            v = self.ask("  > Valor/Detalle (ej. 10 años)")
            items.append({"key": k, "value": v})
            
        section = {
            "title": "Aspectos Técnicos",
            "level": 1,
            "blocks": [
                {"type": "texto", "text_content": desc},
                {"type": "kv_list", "kv_content": items}
            ]
        }
        self.data["sections"].append(section)

    def collect_financial(self):
        self.header("PASO 4: PRESUPUESTO")
        print("Ingrese los ítems de costo principales. (Presione Enter en 'Concepto' para terminar)")
        
        rows = []
        while True:
            concept = self.ask("  > Concepto de Gasto")
            if not concept: break
            val = self.ask("  > Valor Estimado (COP)")
            rows.append([concept, val])
        
        if rows:
            section = {
                "title": "Presupuesto Estimado",
                "level": 1,
                "blocks": [
                    {"type": "texto", "text_content": "A continuación se detalla la estructura de costos directos del proyecto:"},
                    {"type": "tabla", "table_content": {
                        "headers": ["Concepto", "Valor"],
                        "rows": rows,
                        "caption": "Tabla 1. Presupuesto General"
                    }}
                ]
            }
            self.data["sections"].append(section)

    def collect_risks(self):
        self.header("PASO 5: MATRIZ DE RIESGOS")
        print("Identifique los riesgos del proyecto. (Presione Enter en 'Riesgo' para terminar)")
        rows = []
        while True:
            r = self.ask("  > Riesgo Identificado")
            if not r: break
            prob = self.ask("  > Probabilidad (Alta/Media/Baja)")
            mit = self.ask("  > Acción de Mitigación")
            rows.append([r, prob, mit])
            
        if rows:
            section = {
                "title": "Análisis de Riesgos",
                "level": 1,
                "blocks": [
                    {"type": "texto", "text_content": "Se presentan los riesgos previsibles y sus estrategias de manejo:"},
                    {"type": "tabla", "table_content": {
                        "headers": ["Riesgo", "Probabilidad", "Mitigación"],
                        "rows": rows,
                        "caption": "Tabla 2. Matriz de Riesgos"
                    }}
                ]
            }
            self.data["sections"].append(section)

    def generate(self):
        print("\n" + "-"*50)
        print("Generando documento final...")
        try:
            safe_title = self.data.get('project_title', 'Proyecto').replace(' ', '_')[:30]
            if not safe_title: safe_title = "Sin_Titulo"
            
            output_name = f"{safe_title}_MGA_Wizard.docx"
            output_path = os.path.join(os.path.dirname(__file__), "..", "salida", output_name)
            
            # Instanciar Builder en modo MGA
            builder = CITESReportBuilder(output_path, style_mode="MGA")
            
            # Construir Portada
            builder.doc.add_heading(self.data['project_title'], 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p = builder.doc.add_paragraph(self.data.get('entidad', ''))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            
            builder.doc.add_page_break()
            
            # Construir Secciones
            for section in self.data['sections']:
                builder.add_numbered_heading(section['title'], level=section['level'])
                
                for block in section['blocks']:
                    b_type = block['type']
                    if b_type == 'texto':
                         builder.add_text(block.get('text_content'))
                    elif b_type == 'kv_list':
                         builder.add_kv_list(block.get('kv_content', []))
                    elif b_type == 'tabla':
                         builder.add_table(block.get('table_content'))
            
            builder.save()
            print(f"✅ DOCUMENTO CREADO EXITOSAMENTE:")
            print(f"   {output_path}")
            
        except Exception as e:
            print(f"❌ Error generando documento: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    try:
        wizard = MGAWizard()
        wizard.run()
    except KeyboardInterrupt:
        print("\n\nOperación cancelada por el usuario.")
