import streamlit as st
import pandas as pd
import os
import sys
from datetime import date

# Asegurar que podemos importar modulos vecinos
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from software.cites_builder import CITESReportBuilder
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuración de Página
st.set_page_config(page_title="MGA Professional Wizard", layout="wide", page_icon="🧙‍♂️")

# --- INITIALIZATION ---
def init_session():
    if 'project_data' not in st.session_state:
        st.session_state.project_data = {
            "title": "",
            "entity": "",
            "sections": {
                "identificacion": {},
                "tecnica": {},
                "cronograma": [],
                "presupuesto": [],
                "riesgos": []
            } 
        }
    if 'current_step' not in st.session_state:
        st.session_state.current_step = 1

def next_step():
    st.session_state.current_step += 1

def prev_step():
    st.session_state.current_step -= 1

# --- STEPS RENDERERS ---

def render_step_1_metadata():
    st.header("1. Información General del Proyecto")
    st.info("Ingrese los datos básicos que aparecerán en la portada y encabezados.")
    
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.project_data['title'] = st.text_input("Título del Proyecto", st.session_state.project_data['title'], placeholder="Ej. Construcción de Acueducto Veredal...")
    with col2:
        st.session_state.project_data['entity'] = st.text_input("Entidad Proponente", st.session_state.project_data['entity'], placeholder="Ej. Alcaldía Municipal de...")

    st.markdown("---")
    if st.session_state.project_data['title'] and st.session_state.project_data['entity']:
        if st.button("Siguiente: Identificación ➡️"):
            next_step()
            st.rerun()
    else:
        st.warning("Complete los campos para continuar.")

def render_step_2_identification():
    st.header("2. Identificación y Problema")
    st.markdown("Defina claramente el problema que el proyecto busca resolver (Metodología MGA).")
    
    sec = st.session_state.project_data['sections']['identificacion']
    
    sec['problema'] = st.text_area("Descripción del Problema Central", sec.get('problema', ''), height=100)
    
    col1, col2 = st.columns(2)
    with col1:
        sec['poblacion'] = st.text_input("Población Objetivo", sec.get('poblacion', ''))
    with col2:
        sec['ubicacion'] = st.text_input("Ubicación Geográfica", sec.get('ubicacion', ''))
        
    sec['objetivo'] = st.text_area("Objetivo General", sec.get('objetivo', ''), height=100)

    col_back, col_next = st.columns([1, 5])
    with col_back:
        if st.button("⬅️ Atrás"):
            prev_step()
            st.rerun()
    with col_next:
        if sec['problema'] and sec['objetivo']:
            if st.button("Siguiente: Aspectos Técnicos ➡️"):
                next_step()
                st.rerun()
        else:
            st.warning("Problema y Objetivo son obligatorios.")

def render_step_3_technical():
    st.header("3. Aspectos Técnicos")
    st.markdown("Especifique la alternativa técnica seleccionada y sus características.")
    
    sec = st.session_state.project_data['sections']['tecnica']
    
    sec['descripcion'] = st.text_area("Descripción de la Alternativa Técnica", sec.get('descripcion', ''), height=150)
    
    st.subheader("Especificaciones Técnicas")
    st.caption("Agregue detalles como vida útil, dimensiones, normas técnicas, etc.")
    
    if 'specs_df' not in st.session_state:
        # Cargar si existe o crear vacío
        现有_specs = sec.get('especificaciones', [])
        if 现有_specs:
            st.session_state.specs_df = pd.DataFrame(现有_specs)
        else:
            st.session_state.specs_df = pd.DataFrame(columns=["Característica", "Detalle"])

    edited_df = st.data_editor(st.session_state.specs_df, num_rows="dynamic", use_container_width=True)
    sec['especificaciones'] = edited_df.to_dict('records')

    col_back, col_next = st.columns([1, 5])
    with col_back:
        if st.button("⬅️ Atrás", key="back_tech"):
            prev_step()
            st.rerun()
    with col_next:
        if sec['descripcion']:
            if st.button("Siguiente: Cronograma ➡️"):
                next_step()
                st.rerun()
        else:
            st.warning("La descripción técnica es obligatoria.")

def render_step_4_schedule():
    st.header("4. Cronograma de Ejecución")
    st.markdown("Defina las fases y actividades principales del proyecto.")
    
    if 'cronograma_df' not in st.session_state:
        existing = st.session_state.project_data['sections']['cronograma']
        if existing:
            st.session_state.cronograma_df = pd.DataFrame(existing)
        else:
            st.session_state.cronograma_df = pd.DataFrame([
                {"Fase": "Preinversión", "Actividad": "Estudios y Diseños", "Duración (Meses)": "2", "Responsable": "Contratista"},
                {"Fase": "Ejecución", "Actividad": "Obra Civil", "Duración (Meses)": "6", "Responsable": "Contratista"},
            ])

    edited_df = st.data_editor(st.session_state.cronograma_df, num_rows="dynamic", use_container_width=True)
    st.session_state.project_data['sections']['cronograma'] = edited_df.to_dict('records')

    col_back, col_next = st.columns([1, 5])
    with col_back:
        if st.button("⬅️ Atrás", key="back_sched"):
            prev_step()
            st.rerun()
    with col_next:
         if st.button("Siguiente: Presupuesto ➡️"):
            next_step()
            st.rerun()

def render_step_5_budget():
    st.header("5. Presupuesto Detallado")
    st.markdown("Ingrese los costos directos e indirectos del proyecto.")
    
    if 'presupuesto_df' not in st.session_state:
        existing = st.session_state.project_data['sections']['presupuesto']
        if existing:
            st.session_state.presupuesto_df = pd.DataFrame(existing)
        else:
            st.session_state.presupuesto_df = pd.DataFrame([
                {"Ítem": "Personal", "Unidad": "Global", "Cantidad": 1, "Valor Unitario": 5000000},
                {"Ítem": "Materiales", "Unidad": "Global", "Cantidad": 1, "Valor Unitario": 10000000},
            ])

    # Configuración de columnas para cálculos (si fuera más complejo, por ahora simple data entry)
    edited_df = st.data_editor(st.session_state.presupuesto_df, num_rows="dynamic", use_container_width=True)
    
    # Calcular total simple para visualización
    try:
        if not edited_df.empty:
            edited_df['Total'] = edited_df['Cantidad'].astype(float) * edited_df['Valor Unitario'].astype(float)
            grand_total = edited_df['Total'].sum()
            st.metric("Costo Total Estimado", f"${grand_total:,.2f}")
    except:
        st.warning("Asegurese de usar números en Cantidad y Valor Unitario para ver el total.")

    st.session_state.project_data['sections']['presupuesto'] = edited_df.to_dict('records')

    col_back, col_next = st.columns([1, 5])
    with col_back:
        if st.button("⬅️ Atrás", key="back_bud"):
            prev_step()
            st.rerun()
    with col_next:
         if st.button("Siguiente: Riesgos ➡️"):
            next_step()
            st.rerun()

def render_step_6_risks():
    st.header("6. Matriz de Riesgos")
    st.markdown("Identifique riesgos y acciones de mitigación según estándar MGA.")
    
    if 'riesgos_df' not in st.session_state:
        existing = st.session_state.project_data['sections']['riesgos']
        if existing:
            st.session_state.riesgos_df = pd.DataFrame(existing)
        else:
            # Riesgos típicos pre-cargados
            st.session_state.riesgos_df = pd.DataFrame([
                {"Riesgo": "Climáticos", "Probabilidad": "Media", "Impacto": "Alto", "Mitigación": "Programación en tiempo seco"},
                {"Riesgo": "Financieros", "Probabilidad": "Baja", "Impacto": "Alto", "Mitigación": "Aseguramiento de recursos"},
            ])

    edited_df = st.data_editor(st.session_state.riesgos_df, num_rows="dynamic", use_container_width=True, 
                               column_config={
                                   "Probabilidad": st.column_config.SelectboxColumn(options=["Alta", "Media", "Baja"]),
                                   "Impacto": st.column_config.SelectboxColumn(options=["Alto", "Medio", "Bajo"])
                               })
    st.session_state.project_data['sections']['riesgos'] = edited_df.to_dict('records')

    col_back, col_next = st.columns([1, 5])
    with col_back:
        if st.button("⬅️ Atrás", key="back_risk"):
            prev_step()
            st.rerun()
    with col_next:
         if st.button("Finalizar y Generar 🚀", type="primary"):
            next_step()
            st.rerun()

def generate_document():
    st.header("⏳ Generando Documento...")
    try:
        data = st.session_state.project_data
        safe_title = data.get('title', 'Proyecto').replace(' ', '_')[:30]
        output_name = f"{safe_title}_MGA_Pro.docx"
        output_path = os.path.join(os.path.dirname(__file__), "..", "salida", output_name)
        
        # Instanciar Builder
        builder = CITESReportBuilder(output_path, style_mode="MGA_Pro")
        
        # 1. Portada
        builder.doc.add_heading(data['title'], 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = builder.doc.add_paragraph(data['entity'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        builder.doc.add_paragraph(f"Fecha: {date.today()}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        builder.doc.add_page_break()
        
        # 2. Identificación
        sec_id = data['sections']['identificacion']
        builder.add_numbered_heading("Identificación y Objetivos", 1)
        builder.add_text(f"Problema Central: {sec_id.get('problema', '')}")
        builder.add_text(f"Objetivo General: {sec_id.get('objetivo', '')}")
        builder.add_kv_list([
            {"key": "Población Objetivo", "value": sec_id.get('poblacion', '')},
            {"key": "Ubicación", "value": sec_id.get('ubicacion', '')}
        ])
        
        # 3. Técnica
        sec_tec = data['sections']['tecnica']
        builder.add_numbered_heading("Aspectos Técnicos", 1)
        builder.add_text(sec_tec.get('descripcion', ''))
        
        specs = sec_tec.get('especificaciones', [])
        if specs:
            # Convertir lista de dicts a lista de listas para la tabla
            headers = ["Característica", "Detalle"]
            rows = [[item.get("Característica", ""), item.get("Detalle", "")] for item in specs]
            builder.add_table({"headers": headers, "rows": rows, "caption": "Especificaciones Técnicas"})
            
        # 4. Cronograma
        sched = data['sections']['cronograma']
        if sched:
            builder.add_numbered_heading("Cronograma de Ejecución", 1)
            # Asumimos que todas las filas tienen las mismas keys del DF
            if len(sched) > 0:
                headers = list(sched[0].keys())
                rows = [list(item.values()) for item in sched]
                builder.add_table({"headers": headers, "rows": rows, "caption": "Cronograma General"})
                
        # 5. Presupuesto
        bud = data['sections']['presupuesto']
        if bud:
            builder.add_numbered_heading("Presupuesto", 1)
            if len(bud) > 0:
                headers = list(bud[0].keys())
                rows = [list(item.values()) for item in bud]
                builder.add_table({"headers": headers, "rows": rows, "caption": "Presupuesto Detallado"})

        # 6. Riesgos
        risks = data['sections']['riesgos']
        if risks:
            builder.add_numbered_heading("Matriz de Riesgos", 1)
            if len(risks) > 0:
                headers = list(risks[0].keys())
                rows = [list(item.values()) for item in risks]
                builder.add_table({"headers": headers, "rows": rows, "caption": "Matriz de Riesgos"})
        
        builder.save()
        
        st.success("✅ Documento Generado Exitosamente!")
        st.balloons()
        st.markdown(f"**Ubicación:** `{output_path}`")
        
        with open(output_path, "rb") as file:
            st.download_button(
                label="📥 Descargar Documento DOCX",
                data=file,
                file_name=output_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        if st.button("🔄 Crear Nuevo Documento"):
            st.session_state.clear()
            st.rerun()
            
    except Exception as e:
        st.error(f"❌ Error Generando Documento: {e}")
        st.code(e)
        if st.button("⬅️ Regresar"):
            prev_step()
            st.rerun()


from software.ai_engine import analyze_unstructured_text
from software.financial_engine import BudgetManager

def render_step_ai_brain():
    st.header("🧠 Cerebro Artificial (Google Antigravity)")
    st.markdown("""
    **Modo Automático:** Pegue aquí cualquier texto desordenado (notas, correos, ideas, transcripciones) 
    y la IA estructurará todo el proyecto MGA por usted.
    """)
    
    with st.expander("Configuración de API Key (Opcional si ya está en env vars)", expanded=False):
        api_key_input = st.text_input("Gemini API Key", type="password", help="Si no tiene una, se usará el Modo Demo")
    
    user_text = st.text_area("Pegue su información aquí:", height=300, placeholder="Ejemplo: Quiero hacer un parque en el barrio Los Pinos, cuesta como 500 millones, se demora 6 meses...")
    
    if st.button("✨ Estructurar Proyecto Automáticamente", type="primary"):
        with st.spinner("Analizando y estructurando información..."):
            try:
                # Llamada al cerebro
                structured_data = analyze_unstructured_text(user_text, api_key=api_key_input)
                
                if structured_data:
                    st.session_state.project_data = structured_data
                    st.success("¡Proyecto Estructurado! Revise la información en los siguientes pasos.")
                    st.session_state.current_step = 1 # Ir a inicio para revisar
                    st.rerun()
                else:
                    st.error("No se pudo estructurar la información. Intente nuevamente.")
            except Exception as e:
                st.error(f"Error en el proceso de IA: {e}")

def render_step_5_budget():
    st.header("5. Presupuesto Detallado (Motor Financiero)")
    st.markdown("Ingrese los costos. El sistema calculará subtotales y generará el Excel de soporte.")
    
    if 'presupuesto_df' not in st.session_state:
        existing = st.session_state.project_data['sections']['presupuesto']
        if existing:
            st.session_state.presupuesto_df = pd.DataFrame(existing)
        else:
            st.session_state.presupuesto_df = pd.DataFrame([
                {"Ítem": "Personal", "Unidad": "Global", "Cantidad": 1, "Valor Unitario": 5000000},
                {"Ítem": "Materiales", "Unidad": "Global", "Cantidad": 1, "Valor Unitario": 10000000},
            ])

    # Editor de Datos
    edited_df = st.data_editor(st.session_state.presupuesto_df, num_rows="dynamic", use_container_width=True)
    st.session_state.project_data['sections']['presupuesto'] = edited_df.to_dict('records') # Guardar estado raw

    # --- CÁLCULO EN TIEMPO REAL USANDO BUDGET MANAGER ---
    bm = BudgetManager(edited_df)
    total = bm.calculate_totals()
    
    col_metric1, col_metric2 = st.columns(2)
    col_metric1.metric("Costo Directo Total", f"${total:,.0f}")
    col_metric2.metric("IVA Estimado (19%)", f"${total * 0.19:,.0f}")

    col_back, col_next = st.columns([1, 5])
    with col_back:
        if st.button("⬅️ Atrás", key="back_bud"):
            prev_step()
            st.rerun()
    with col_next:
         if st.button("Siguiente: Riesgos ➡️"):
            next_step()
            st.rerun()

def generate_document():
    st.header("⏳ Generando Entregables...")
    try:
        data = st.session_state.project_data
        safe_title = data.get('title', 'Proyecto').replace(' ', '_')[:30]
        
        # Rutas de Salida
        docx_name = f"{safe_title}_MGA_Pro.docx"
        xlsx_name = f"{safe_title}_Presupuesto.xlsx"
        
        output_dir = os.path.join(os.path.dirname(__file__), "..", "salida")
        os.makedirs(output_dir, exist_ok=True)
        
        docx_path = os.path.join(output_dir, docx_name)
        xlsx_path = os.path.join(output_dir, xlsx_name)
        
        # --- 1. GENERAR EXCEL (Financiero) ---
        bud_data = data['sections']['presupuesto']
        bm = BudgetManager(bud_data)
        bm.export_excel_summary(xlsx_path)
        
        # --- 2. GENERAR DOCX (Técnico) ---
        builder = CITESReportBuilder(docx_path, style_mode="MGA_Pro")
        
        # Portada
        builder.doc.add_heading(data['title'], 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = builder.doc.add_paragraph(data['entity'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        builder.doc.add_paragraph(f"Fecha: {date.today()}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        builder.doc.add_page_break()
        
        # Identificación
        sec_id = data['sections']['identificacion']
        builder.add_numbered_heading("Identificación y Objetivos", 1)
        builder.add_text(f"Problema Central: {sec_id.get('problema', '')}")
        builder.add_text(f"Objetivo General: {sec_id.get('objetivo', '')}")
        builder.add_kv_list([
            {"key": "Población Objetivo", "value": sec_id.get('poblacion', '')},
            {"key": "Ubicación", "value": sec_id.get('ubicacion', '')}
        ])
        
        # Técnica
        sec_tec = data['sections']['tecnica']
        builder.add_numbered_heading("Aspectos Técnicos", 1)
        builder.add_text(sec_tec.get('descripcion', ''))
        specs = sec_tec.get('especificaciones', [])
        if specs:
            headers = ["Característica", "Detalle"]
            rows = [[item.get("Característica", ""), item.get("Detalle", "")] for item in specs]
            builder.add_table({"headers": headers, "rows": rows, "caption": "Especificaciones Técnicas"})
            
        # Cronograma
        sched = data['sections']['cronograma']
        if sched:
            builder.add_numbered_heading("Cronograma de Ejecución", 1)
            if len(sched) > 0:
                headers = list(sched[0].keys())
                rows = [list(item.values()) for item in sched]
                builder.add_table({"headers": headers, "rows": rows, "caption": "Cronograma General"})
                
        # Presupuesto (Usando tabla resumen formateada del BudgetManager)
        if bud_data:
            builder.add_numbered_heading("Presupuesto Resumido", 1)
            summary_table = bm.get_summary_table_for_doc()
            if summary_table:
                try:
                    headers = list(summary_table[0].keys())
                    rows = [list(item.values()) for item in summary_table]
                    builder.add_table({"headers": headers, "rows": rows, "caption": "Tabla Resumen de Costos (Ver Anexo Excel)"})
                except:
                    builder.add_text("Ver detalle en anexo Excel.")

        # Riesgos
        risks = data['sections']['riesgos']
        if risks:
            builder.add_numbered_heading("Matriz de Riesgos", 1)
            if len(risks) > 0:
                headers = list(risks[0].keys())
                rows = [list(item.values()) for item in risks]
                builder.add_table({"headers": headers, "rows": rows, "caption": "Matriz de Riesgos"})
        
        builder.save()
        
        st.success("✅ ¡Paquete de Proyecto Generado!")
        st.balloons()
        
        col_d1, col_d2 = st.columns(2)
        
        with open(docx_path, "rb") as f_docx:
            col_d1.download_button(
                label="📄 Descargar Informe (.docx)",
                data=f_docx,
                file_name=docx_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        with open(xlsx_path, "rb") as f_xlsx:
            col_d2.download_button(
                label="📊 Descargar Cálculos (.xlsx)",
                data=f_xlsx,
                file_name=xlsx_name,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            
        if st.button("🔄 Crear Nuevo Proyecto"):
            st.session_state.clear()
            st.rerun()
            
    except Exception as e:
        st.error(f"❌ Error Generando Documentos: {e}")
        st.code(e)
        import traceback
        st.code(traceback.format_exc())
        if st.button("⬅️ Regresar"):
            prev_step()
            st.rerun()

    init_session()
    
    # Sidebar
    st.sidebar.title("Navegación")
    
    # Modo Directo en Sidebar
    if st.sidebar.button("🧠 MODO INTELIGENTE"):
        st.session_state.current_step = 0 # 0 será el paso de IA
        st.rerun()

    # Progress Bar
    steps = ["Cerebro IA", "Inicio", "Identificación", "Técnica", "Cronograma", "Presupuesto", "Riesgos", "Generar"]
    # Ajuste de índice para progress bar (step 0 es index 0, step 1 es index 1...)
    curr = st.session_state.current_step
    
    # Normalizar progress bar
    # Si curr = 0 (IA), progress = 5%
    if curr == 0:
        st.progress(5)
        st.caption("Modo Inteligente: Ingesta de Datos")
        render_step_ai_brain()
    else:
        # Steps 1 a 7
        progress = int((curr / 7) * 100)
        st.progress(progress)
        st.caption(f"Paso {curr} de 7: {steps[curr]}")
        
        if curr == 1:
            render_step_1_metadata()
        elif curr == 2:
            render_step_2_identification()
        elif curr == 3:
            render_step_3_technical()
        elif curr == 4:
            render_step_4_schedule()
        elif curr == 5:
            render_step_5_budget()
        elif curr == 6:
            render_step_6_risks()
        elif curr == 7:
            generate_document()

if __name__ == "__main__":
    main()

