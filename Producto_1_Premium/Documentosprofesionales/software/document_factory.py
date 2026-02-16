from typing import List, Dict, Any
# Asumimos que python-docx está instalado, si no, se debe instalar.
try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Advertencia: python-docx no está instalado. Instalelo con 'pip install python-docx'")
    Document = None

class DocumentFactory:
    """Fábrica de Documentos Profesionales (MGA + APA)."""
    
    def __init__(self, output_path: str):
        self.output_path = output_path
        self.doc = Document()
        self._configurar_estilos()

    def _configurar_estilos(self):
        """Configura márgenes y fuentes APA."""
        section = self.doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Párrafo
        paragraph_format = style.paragraph_format
        # APA requiere doble espacio, pero en reportes técnicos MGA se usa a veces 1.5
        paragraph_format.line_spacing = 1.5 

    def agregar_titulo(self, texto: str, nivel: int = 1):
        self.doc.add_heading(texto, level=nivel)

    def agregar_parrafo(self, texto: str):
        p = self.doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def agregar_lista(self, items: List[str], titulo: str = None):
        if titulo:
            p = self.doc.add_paragraph(titulo)
            p.bold = True
        for item in items:
            self.doc.add_paragraph(item, style='List Bullet')

    def procesar_contenido(self, contenido: Dict[str, Any]):
        """Procesa un diccionario de contenido estructurado (Renderizado por Módulos MGA)."""
        if "titulo" in contenido:
            self.agregar_titulo(contenido["titulo"], 1)
        
        if "cuerpo" in contenido:
            for elemento in contenido["cuerpo"]:
                tipo = elemento.get("tipo")
                texto = elemento.get("texto", "")
                
                if tipo == "parrafo":
                    self.agregar_parrafo(texto)
                elif tipo == "titulo2":
                    self.agregar_titulo(texto, 2)
                elif tipo == "lista":
                    self.agregar_lista(elemento.get("items", []), elemento.get("titulo"))
                elif tipo == "tabla":
                    # Lógica simple de tabla
                    datos = elemento.get("datos", [])
                    if datos:
                        table = self.doc.add_table(rows=1, cols=len(datos[0]))
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        headers = list(datos[0].keys())
                        for i, h in enumerate(headers):
                            hdr_cells[i].text = h
                        
                        for row_data in datos:
                            row_cells = table.add_row().cells
                            for i, h in enumerate(headers):
                                row_cells[i].text = str(row_data[h])

    def guardar(self):
        self.doc.save(self.output_path)
        print(f"Documento guardado exitosamente en: {self.output_path}")

# Ejemplo de orquestación
class ProjectAssembler:
    """Ensamblador del Proyecto Completo."""
    
    def __init__(self, filename: str):
        self.factory = DocumentFactory(filename)
        self.modulos = []

    def registrar_modulo(self, modulo):
        self.modulos.append(modulo)

    def construir(self):
        for modulo in self.modulos:
            contenido = modulo.render_content()
            self.factory.procesar_contenido(contenido)
            self.factory.doc.add_page_break()
        self.factory.guardar()
